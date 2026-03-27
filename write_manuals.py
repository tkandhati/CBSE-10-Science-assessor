from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    return p

def add_normal(doc, text, bold=False, centered=False):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.bold = bold
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_bold_inline(doc, parts):
    """parts = list of (text, bold) tuples"""
    p = doc.add_paragraph(style='List Bullet')
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — Student Manual
# ══════════════════════════════════════════════════════════════════════════════

doc1_path = r'C:\CBSE10\science\Documents\Technical Documents\Science_Assessor_Student_Manual_v1_0.docx'

doc = Document()

# Title block — Normal, centered
p = doc.add_paragraph(style='Normal')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Science Assessor — Student Manual')
run.bold = True

p = doc.add_paragraph(style='Normal')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('CBSE Class 10 Science Assessment Platform')

p = doc.add_paragraph(style='Normal')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Version 1.0')

# 1. Overview
add_heading(doc, '1. Overview', 1)
doc.add_paragraph(
    'Science Assessor covers all 13 CBSE Class 10 Science chapters across Chemistry, Biology, Physics, and Environmental Science. Three session types are available:',
    style='Normal'
)
add_bold_inline(doc, [('Understanding Session', True), (' — 10–12 questions in the browser, auto-scored instantly', False)])
add_bold_inline(doc, [('Chapter Test', True), (' — exam-pattern paper, write on paper, upload PDF for AI scoring', False)])
add_bold_inline(doc, [('Full Mock Test', True), (' — complete CBSE paper (84 marks, 3 hours), same upload process', False)])

# 2. Dashboard
add_heading(doc, '2. Dashboard', 1)
doc.add_paragraph('Your dashboard shows:', style='Normal')
add_bullet(doc, 'XP earned and current level')
add_bullet(doc, 'Streak — consecutive days with at least one session (resets if you miss a day)')
add_bullet(doc, 'Chapter performance bars — score per chapter, click to expand topic breakdown')
add_bullet(doc, 'Badges earned')

# 3. Understanding Sessions
add_heading(doc, '3. Understanding Sessions', 1)
doc.add_paragraph('Daily practice. Questions are selected to target your weak topics automatically.', style='Normal')

add_heading(doc, '3.1 Starting', 2)
add_bold_inline(doc, [('Tap ', False), ('Start Session → Understanding Session', True)])
add_bullet(doc, 'Select a chapter; optionally narrow to a topic')
add_bold_inline(doc, [('Tap ', False), ('Start', True)])

add_heading(doc, '3.2 Answering', 2)
add_bullet(doc, 'MCQ: tap your choice — scored immediately')
add_bold_inline(doc, [('Short answer / numerical: type in the text box, tap ', False), ('Next', True)])
add_bullet(doc, 'Timer runs per question for tracking only — no cut-off')

add_heading(doc, '3.3 Numerical questions', 2)
doc.add_paragraph('The numbers change each time after a correct answer to prevent memorisation.', style='Normal')

add_heading(doc, '3.4 Results', 2)
add_bullet(doc, 'Green — correct; Amber — partial marks; Red — incorrect')
add_bullet(doc, 'Model answer shown for every question')

# 4. Chapter Tests
add_heading(doc, '4. Chapter Tests', 1)
doc.add_paragraph('Exam-pattern test for one chapter. Write answers on paper, upload a single PDF.', style='Normal')

add_heading(doc, '4.1 Before you start', 2)
add_bullet(doc, 'Have paper, pen, and a phone/scanner ready to photograph your answers')

add_heading(doc, '4.2 Starting', 2)
add_bold_inline(doc, [('Tap ', False), ('Start Session → Chapter Test', True)])
add_bullet(doc, 'Choose Short (15 marks, 20 min) or Regular (40 marks, 45 min)')
add_bold_inline(doc, [('Select chapter → ', False), ('Generate Test', True)])
add_bold_inline(doc, [('Write answers on paper, then tap ', False), ('Mark Done Writing', True)])

add_heading(doc, '4.3 Uploading your answer sheet', 2)
add_bullet(doc, 'Photograph all answer pages and combine into a single PDF (under 20 MB)')
add_bold_inline(doc, [('Tap ', False), ('Upload Answers', True), (', select the PDF, tap ', False), ('Submit', True)])

add_heading(doc, '4.4 OCR confirmation', 2)
doc.add_paragraph(
    'The app extracts handwritten text automatically. If confidence is low it shows the extracted text — confirm it is correct or edit before scoring proceeds.',
    style='Normal'
)

add_heading(doc, '4.5 Results', 2)
add_bullet(doc, 'Total score and percentage')
add_bullet(doc, 'Per-question feedback (green / amber / red) with model answers')
add_bullet(doc, 'Improvement suggestions per question')
add_bullet(doc, 'Marks-lost summary by question type')

# 5. Full Mock Test
add_heading(doc, '5. Full Mock Test', 1)
doc.add_paragraph('Complete CBSE Class 10 Science paper — all 13 chapters, 84 marks, 3 hours.', style='Normal')
add_bold_inline(doc, [('Tap ', False), ('Start Session → Full Mock Test → Start Mock', True)])
add_bullet(doc, 'Write all answers on paper; set your own 3-hour timer')
add_bold_inline(doc, [('Tap ', False), ('Mark Done Writing', True), (', upload a single PDF of all pages', False)])

doc.add_paragraph('Results include:', style='Normal')
add_bullet(doc, 'Section-wise and chapter-wise score breakdown')
add_bullet(doc, 'Exam readiness estimate out of 84 with subject breakdown')

# 6. XP, Streaks, and Badges
add_heading(doc, '6. XP, Streaks, and Badges', 1)
add_bullet(doc, 'Every answer earns XP based on difficulty; correct = full XP, partial = proportional')
add_bullet(doc, 'Every 500 XP = level up')
add_bullet(doc, 'Maintain a daily session to keep your streak alive')
add_bullet(doc, 'Chapter Master badges unlock when you score ≥ 80% on a chapter across 5+ sessions')

# 7. OCR Troubleshooting (table)
add_heading(doc, '7. OCR Troubleshooting', 1)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Problem'
hdr[1].text = 'Fix'

rows_data = [
    ('Low confidence warning', 'Review extracted text, edit if wrong, tap Confirm'),
    ('Handwriting not recognised', 'Write larger, use dark ink, photograph in good light'),
    ('Diagram not read', 'Describe the diagram in words in the correction box'),
    ('Upload fails', 'Check PDF is under 20 MB; reduce photo resolution'),
]
for i, (problem, fix) in enumerate(rows_data, start=1):
    table.rows[i].cells[0].text = problem
    table.rows[i].cells[1].text = fix

doc.save(doc1_path)
print(f'Saved: {doc1_path}')


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 — Admin Manual
# ══════════════════════════════════════════════════════════════════════════════

doc2_path = r'C:\CBSE10\science\Documents\Technical Documents\Science_Assessor_Admin_Manual_v1_0.docx'

doc2 = Document()

# Title block — Normal (not centered per spec)
p = doc2.add_paragraph(style='Normal')
run = p.add_run('Science Assessor — Admin Manual')
run.bold = True

doc2.add_paragraph('CBSE Class 10 Science Assessment Platform', style='Normal')
doc2.add_paragraph('Version 1.0', style='Normal')

# 1. Overview
add_heading(doc2, '1. Overview', 1)
p = doc2.add_paragraph(style='Normal')
p.add_run('The Admin view monitors student progress, manages the question bank, and reviews AI-scored sessions. Access it via the ')
r = p.add_run('Admin')
r.bold = True
p.add_run(' tab.')

# 2. Admin Dashboard
add_heading(doc2, '2. Admin Dashboard  (/admin)', 1)

add_heading(doc2, 'Summary strip (top of page):', 2)
add_bullet(doc2, 'Total sessions, total questions answered, overall average score, current streak, best streak')

add_heading(doc2, 'Chapter performance bars:', 2)
add_bullet(doc2, 'One bar per chapter, colour-coded by score')
add_bullet(doc2, 'Click any bar to expand topic-level scores within that chapter')

add_heading(doc2, 'Strengths and Weaknesses panels:', 2)
add_bullet(doc2, 'Strengths: top 4 topics by score (minimum 5 attempts) — green pills')
add_bullet(doc2, 'Weaknesses: top 4 topics by score (minimum 3 attempts) — red/orange pills')

add_heading(doc2, 'Exam Readiness Estimate:', 2)
doc2.add_paragraph(
    'Projected board score out of 84 based on weighted topic performance:',
    style='Normal'
)
add_bullet(doc2, 'Physics 27 marks, Chemistry 27 marks, Biology 27 marks, Environmental Science 3 marks')

add_heading(doc2, 'Coverage Gap Alert:', 2)
doc2.add_paragraph(
    'Appears when the question bank is too thin for a topic to generate a balanced test. Click to go directly to the Question Bank.',
    style='Normal'
)

# 3. Topic Intelligence
add_heading(doc2, '3. Topic Intelligence  (/admin/strengths)', 1)
add_bold_inline(doc2, [('Weak Topics panel', True), (' — 5 weakest topics across all chapters, ranked by score', False)])
add_bold_inline(doc2, [('Untested Topics panel', True), (' — topics with fewer than 3 questions answered (blind spots)', False)])
add_bold_inline(doc2, [('Chapter accordion', True), (' — all 13 chapters grouped by subject; each topic row shows score, attempts, and last attempted date', False)])

# 4. Session History
add_heading(doc2, '4. Session History  (/admin/sessions and /admin/session/:id)', 1)

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Session list:')
r.bold = True
p.add_run(' All sessions in reverse order. Filter by type or chapter. Columns: date, type, chapter, score, percentage.')

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Session detail (click any row):')
r.bold = True

add_bullet(doc2, 'Every question with student\'s answer (typed or OCR-extracted)')
add_bullet(doc2, 'Score per question, evaluation method (Deterministic / Keyword / AI)')
add_bullet(doc2, 'Green/amber/red feedback and model answer')
add_bullet(doc2, 'AI improvement suggestions per question')
add_bullet(doc2, 'Marks-lost analysis by type and reason')

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Score override:')
r.bold = True

add_bold_inline(doc2, [('Click ', False), ('Override', True), (' on any AI-scored answer', False)])
add_bullet(doc2, 'Enter corrected score and reason, click Save — session total recalculates immediately')

# 5. Study Guidance
add_heading(doc2, '5. Study Guidance  (/admin/guidance)', 1)
doc2.add_paragraph('AI-generated weekly study plan, refreshed every 24 hours. Contains:', style='Normal')
add_bullet(doc2, '3 priority topics for the week with NCERT section references')
add_bullet(doc2, 'Day-by-day session plan for the next 7 days')
add_bullet(doc2, 'Exam readiness projection with impact of improving specific weak topics')

p = doc2.add_paragraph(style='Normal')
p.add_run('To force a refresh click ')
r = p.add_run('Refresh Now')
r.bold = True
p.add_run('.')

# 6. Question Bank
add_heading(doc2, '6. Question Bank  (/admin/questions)', 1)

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Stats panel:')
r.bold = True
p.add_run(' Approved questions by chapter, type (MCQ / short / numerical / long / diagram / assertion-reason / case-based), and difficulty (L1–L5).')

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Coverage report:')
r.bold = True
p.add_run(' Per-chapter health — green (sufficient), amber (some topics thin), red (insufficient for test generation).')

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Adding questions — PDF upload:')
r.bold = True

# numbered steps as bullets
p = doc2.add_paragraph(style='List Bullet')
p.add_run('Drag a CBSE paper or reference PDF onto the upload area')

p = doc2.add_paragraph(style='List Bullet')
p.add_run('Tap ')
r = p.add_run('Upload and Scan')
r.bold = True
p.add_run(' — AI extracts and tags questions (1–3 min)')

p = doc2.add_paragraph(style='List Bullet')
p.add_run('Go to ')
r = p.add_run('Review Queue')
r.bold = True
p.add_run(' to approve or reject each extracted question')

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Review Queue actions per question:')
r.bold = True

add_bold_inline(doc2, [('Approve', True), (' — goes live immediately', False)])
add_bold_inline(doc2, [('Edit then Approve', True), (' — adjust chapter, topic, type, difficulty, marks, or question text', False)])
add_bold_inline(doc2, [('Reject', True), (' — discards the question', False)])

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Live Bank tab:')
r.bold = True
p.add_run(' All approved questions. Filter by chapter, topic, type, or difficulty. Edit or remove individual questions.')

p = doc2.add_paragraph(style='Normal')
r = p.add_run('Re-tagging:')
r.bold = True
p.add_run(' Select a chapter/topic and tap ')
r = p.add_run('Run Retag')
r.bold = True
p.add_run(' to have AI update metadata. Updated questions return to Review Queue before going live.')

# 7. System Notes
add_heading(doc2, '7. System Notes', 1)
add_bold_inline(doc2, [('Chapter 14 (Sources of Energy)', True), (' — not in CBSE 2026 syllabus; no questions in the bank', False)])
add_bold_inline(doc2, [('Healthy bank minimums:', True), (' 20 approved questions per chapter; at least 3 questions per topic-type combination; at least 5 long-answer and 3 case-based per chapter for Full Mock', False)])
add_bold_inline(doc2, [('Coverage gap alert:', True), (' click the alert → add questions via PDF upload → approve → alert clears on next test generation', False)])

doc2.save(doc2_path)
print(f'Saved: {doc2_path}')
